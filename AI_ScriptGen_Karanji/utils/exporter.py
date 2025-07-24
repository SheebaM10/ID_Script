from docx import Document
from docx.shared import Pt
import re

def add_markdown_text(cell, text):
    # Replace <...> with _..._ for consistent handling
    text = re.sub(r'<([^>]+)>', r'_\1_', text)
    # Split by bold (**...**) and italic (_..._)
    pattern = r'(\*\*.*?\*\*|_[^_]+_|<[^>]+>)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = cell.paragraphs[0].add_run(part[2:-2])
            run.bold = True
        elif (part.startswith('_') and part.endswith('_')):
            run = cell.paragraphs[0].add_run(part[1:-1])
            run.italic = True
        else:
            cell.paragraphs[0].add_run(part)

def export_to_docx_table(rows, output_path, title="Instructional Script"):
    doc = Document()
    doc.add_heading(title, 0)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Slide Number"
    hdr_cells[1].text = "Voice-Over Script"
    hdr_cells[2].text = "On-Screen Text"
    hdr_cells[3].text = "Video Description"
    hdr_cells[4].text = "Image/Infographic Suggestion"

    for row in rows:
        row_cells = table.add_row().cells
        for i, key in enumerate(["slide_number", "voice_over", "on_screen_text", "video_description", "image_suggestion"]):
            cell = row_cells[i]
            cell.text = ""  # Clear default paragraph
            add_markdown_text(cell, row[key].replace('\\n', '\n'))

    doc.save(output_path)
